import os
import io
import json
import pdfplumber
import anthropic
import httpx
from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
_http_client = httpx.Client(
    timeout=httpx.Timeout(connect=10.0, read=75.0, write=10.0, pool=5.0)
)
client = anthropic.Anthropic(
    api_key=os.environ.get('ANTHROPIC_API_KEY'),
    http_client=_http_client
)

DOMAINS = [
    "Domain 1: Culture",
    "Domain 2: Systems & Processes",
    "Domain 3: Workforce Development",
    "Domain 4: Service Delivery",
    "Domain 5: Governance, Leadership and Management & Quality",
    "Domain 6: Environment"
]

DOMAIN_TO_FULL_TABLE = {
    "Domain 1": 2, "Domain 2": 3, "Domain 3": 4,
    "Domain 4": 5, "Domain 5": 6, "Domain 6": 7
}

DOMAIN_TO_SUMMARY_TABLE = {
    "Domain 1": 8, "Domain 2": 9, "Domain 3": 10,
    "Domain 4": 11, "Domain 5": 12, "Domain 6": 13
}

INSUFFICIENT = "INSUFFICIENT_INFO"


def extract_pdf_text(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
    return text


def get_domain_key(domain_name):
    for key in DOMAIN_TO_FULL_TABLE.keys():
        if key in domain_name:
            return key
    return None


def clear_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    for i in range(len(cell.paragraphs) - 1, 0, -1):
        p = cell.paragraphs[i]._element
        p.getparent().remove(p)


def set_cell_italic(cell, text):
    clear_cell(cell)
    para = cell.paragraphs[0]
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    for i, p_text in enumerate(paragraphs):
        if i == 0:
            run = para.add_run(p_text)
        else:
            new_para = cell.add_paragraph()
            run = new_para.add_run(p_text)
        run.italic = True
        run.font.size = Pt(10)


def set_cell_normal(cell, text, bold_first_line=False):
    clear_cell(cell)
    para = cell.paragraphs[0]
    paragraphs = [p for p in text.split('\n') if p.strip()]
    for i, p_text in enumerate(paragraphs):
        if i == 0:
            run = para.add_run(p_text)
            run.bold = bold_first_line
        else:
            new_para = cell.add_paragraph()
            run = new_para.add_run(p_text)
        run.font.size = Pt(10)


def set_cell_insufficient(cell):
    """Set bold red warning text when transcript doesn't have enough info."""
    clear_cell(cell)
    para = cell.paragraphs[0]
    run = para.add_run("INSUFFICIENT INFORMATION IN TRANSCRIPT TO COMPLETE THIS SECTION — PLEASE COMPLETE MANUALLY")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def maybe_set_cell(cell, text, italic=False):
    """Set cell content, or insufficient warning if text is INSUFFICIENT_INFO."""
    if not text or text.strip() == INSUFFICIENT:
        set_cell_insufficient(cell)
    elif italic:
        set_cell_italic(cell, text)
    else:
        set_cell_normal(cell, text)


def set_tick(cell, standard_met):
    """Keep only the correct tick/cross in a Standard Met cell.
    Para 0 = green tick (achieved)
    Para 1 = red cross (not_achieved)
    Para 2 = orange tick (partial)
    """
    status_to_keep = {"achieved": 0, "not_achieved": 1, "partial": 2}
    keep_idx = status_to_keep.get(standard_met, 0)
    paras = cell.paragraphs
    if len(paras) >= 3:
        for i in range(len(paras) - 1, -1, -1):
            if i != keep_idx:
                p = paras[i]._element
                p.getparent().remove(p)


def replace_paragraph_text(doc, search_text, new_text):
    """Find a paragraph containing search_text and replace its full text."""
    for para in doc.paragraphs:
        if search_text in para.text:
            # Clear all runs
            for run in para.runs:
                run.text = ""
            # Set text in first run, or add one
            if para.runs:
                para.runs[0].text = new_text
            else:
                run = para.add_run(new_text)
            return True
    return False


def get_claude_content(pdf_text, transcript, focus_domain_1, focus_domain_2):
    prompt = f"""You are helping complete a Headway Approved Provider Interim Review Report.

FOCUS DOMAINS FOR THIS REVIEW:
- Focus Domain 1: {focus_domain_1}
- Focus Domain 2: {focus_domain_2}

Return ONLY a valid JSON object. Use "INSUFFICIENT_INFO" as the value for any field where there is not enough information in the reaccreditation report or transcript to complete it accurately.

{{
  "unit_name": "name of the unit/service from reaccreditation report, or INSUFFICIENT_INFO",
  "company_name": "name of the provider company, or INSUFFICIENT_INFO",
  "town_city": "town or city of the unit, or INSUFFICIENT_INFO",
  "reaccreditation_date": "date of reaccreditation",
  "overall_grade": "Outstanding/Good/Adequate/Inadequate",
  "summary_scores": [
    {{"domain": "Domain 1: Culture", "rating": "Good"}},
    {{"domain": "Domain 2: Systems & Processes", "rating": "Outstanding"}},
    {{"domain": "Domain 3: Workforce Development", "rating": "Outstanding"}},
    {{"domain": "Domain 4: Service Delivery", "rating": "Outstanding"}},
    {{"domain": "Domain 5: Governance, Leadership and Management & Quality", "rating": "Good"}},
    {{"domain": "Domain 6: Environment", "rating": "Good"}}
  ],
  "focus_domains": [
    {{
      "domain_name": "{focus_domain_1}",
      "standards": [
        {{
          "standard_number": "AP 1",
          "standard_name": "Standard name",
          "evidence_text": "Copy EXACT verbatim text from the reaccreditation report for this standard, or INSUFFICIENT_INFO",
          "standard_met": "achieved, partial, or not_achieved — based on the reaccreditation report evidence"
        }}
      ],
      "evaluation_and_comment": "3-5 paragraphs from transcript for this domain, or INSUFFICIENT_INFO if transcript does not cover it",
      "qi_recommendations": [
        {{"title": "Title", "full_text": "Full verbatim text from reaccreditation report, or INSUFFICIENT_INFO"}}
      ],
      "qi_recommendation_updates": [
        {{"title": "Title", "update_text": "Update from transcript, or INSUFFICIENT_INFO"}}
      ],
      "qi_suggestions": [
        {{"title": "Title", "full_text": "Full verbatim text from reaccreditation report, or INSUFFICIENT_INFO"}}
      ],
      "qi_suggestion_updates": [
        {{"title": "Title", "update_text": "Update from transcript, or INSUFFICIENT_INFO"}}
      ],
      "outcome": "PASS",
      "rating": "Good"
    }},
    {{
      "domain_name": "{focus_domain_2}",
      "standards": [],
      "evaluation_and_comment": "INSUFFICIENT_INFO",
      "qi_recommendations": [],
      "qi_recommendation_updates": [],
      "qi_suggestions": [],
      "qi_suggestion_updates": [],
      "outcome": "PASS",
      "rating": "Good"
    }}
  ],
  "other_domains": [
    {{
      "domain_name": "Domain X: Name",
      "evaluation_and_comment": "1-2 paragraphs from transcript, or INSUFFICIENT_INFO",
      "qi_recommendations": [{{"title": "Title", "full_text": "Text or INSUFFICIENT_INFO"}}],
      "qi_recommendation_updates": [{{"title": "Title", "update_text": "Text or INSUFFICIENT_INFO"}}],
      "qi_suggestions": [{{"title": "Title", "full_text": "Text or INSUFFICIENT_INFO"}}],
      "qi_suggestion_updates": [{{"title": "Title", "update_text": "Text or INSUFFICIENT_INFO"}}],
      "outcome": "PASS",
      "rating": "Good"
    }}
  ]
}}

REACCREDITATION REPORT:
{pdf_text[:10000]}

INTERIM REVIEW TRANSCRIPT:
{transcript[:4000]}
"""

    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = message.content[0].text.strip()
    if response_text.startswith("```"):
        response_text = response_text.split("```")[1]
        if response_text.startswith("json"):
            response_text = response_text[4:]
        response_text = response_text.rsplit("```", 1)[0]

    return json.loads(response_text.strip())


def populate_document(template_path, data):
    doc = Document(template_path)
    tables = doc.tables

    # --- Header: Unit, Company, Town/City ---
    unit = data.get("unit_name", "")
    company = data.get("company_name", "")
    town = data.get("town_city", "")
    parts = [p for p in [unit, company, town] if p and p != INSUFFICIENT]
    header_text = ", ".join(parts) if parts else "INSUFFICIENT INFORMATION — PLEASE COMPLETE MANUALLY"
    replace_paragraph_text(doc, "Unit, Company, Town/City", header_text)

    # --- Summary scores table (Table 0) ---
    scores_table = tables[0]
    for row in scores_table.rows[2:8]:
        domain_text = row.cells[0].text.strip()
        for score in data.get("summary_scores", []):
            if score["domain"].lower() in domain_text.lower() or domain_text.lower() in score["domain"].lower():
                set_cell_normal(row.cells[1], score["rating"].upper())
                break

    # --- Focus domain full tables ---
    for fd in data.get("focus_domains", []):
        dkey = get_domain_key(fd["domain_name"])
        if not dkey:
            continue
        tidx = DOMAIN_TO_FULL_TABLE.get(dkey)
        if tidx is None or tidx >= len(tables):
            continue

        table = tables[tidx]
        standards = fd.get("standards", [])
        std_idx = 0

        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if "paste from most recent" in txt.lower() and std_idx < len(standards):
                    ev = standards[std_idx].get("evidence_text", INSUFFICIENT)
                    status = standards[std_idx].get("standard_met", "achieved")
                    maybe_set_cell(cell, ev, italic=True)
                    # Set correct tick in Standard Met column (col 3)
                    if len(row.cells) > 3:
                        set_tick(row.cells[3], status)
                    std_idx += 1
                    break

        # Evaluation and comment
        for row in table.rows:
            for cell in row.cells:
                if "Evaluation and Comment" in cell.text and ("If Outstanding" in cell.text or "All standards" in cell.text):
                    ev = fd.get("evaluation_and_comment", INSUFFICIENT)
                    if ev == INSUFFICIENT:
                        set_cell_insufficient(cell)
                    else:
                        set_cell_normal(cell, "Evaluation and Comment:\n\n" + ev)
                    break

        # QI tables
        qi_recs = fd.get("qi_recommendations", [])
        qi_rec_upd = fd.get("qi_recommendation_updates", [])
        qi_sugs = fd.get("qi_suggestions", [])
        qi_sug_upd = fd.get("qi_suggestion_updates", [])

        in_rec = False
        in_sug = False
        rec_i = 0
        sug_i = 0

        for row in table.rows:
            row_text = " ".join(c.text for c in row.cells)

            if "Quality Improvement Recommendations at" in row_text:
                in_rec = True; in_sug = False; continue
            if "Quality Improvement Suggestions at" in row_text:
                in_sug = True; in_rec = False; continue
            if "Outcome:" in row_text:
                in_rec = False; in_sug = False
                for cell in row.cells:
                    if "Outcome:" in cell.text:
                        set_cell_normal(cell, f"Outcome: {fd.get('outcome', 'PASS')}")
                    elif "Rating" in cell.text:
                        set_cell_normal(cell, f"Reaccreditation Assessment Rating: {fd.get('rating', 'Good')}\nInterim Review Rating: {fd.get('rating', 'Good')}")
                continue

            if in_rec and len(row.cells) >= 2:
                lc = row.cells[0]
                rc = row.cells[-1]
                if lc.text.strip() in ["", "Add a title in bold", "1.", "2.", "3."] or "Add a title" in lc.text:
                    if rec_i < len(qi_recs):
                        r = qi_recs[rec_i]
                        ft = r.get("full_text", INSUFFICIENT)
                        if ft == INSUFFICIENT:
                            set_cell_insufficient(lc)
                        else:
                            set_cell_normal(lc, f"{rec_i + 1}. {r['title']}:\n{ft}")
                        if rec_i < len(qi_rec_upd):
                            u = qi_rec_upd[rec_i]
                            ut = u.get("update_text", INSUFFICIENT)
                            if ut == INSUFFICIENT:
                                set_cell_insufficient(rc)
                            else:
                                set_cell_normal(rc, f"{rec_i + 1}. {u['title']}:\n{ut}")
                        rec_i += 1

            if in_sug and len(row.cells) >= 2:
                lc = row.cells[0]
                rc = row.cells[-1]
                if lc.text.strip() in ["", "Add a title in bold", "1.", "2.", "3."] or "Add a title" in lc.text:
                    if sug_i < len(qi_sugs):
                        s = qi_sugs[sug_i]
                        ft = s.get("full_text", INSUFFICIENT)
                        if ft == INSUFFICIENT:
                            set_cell_insufficient(lc)
                        else:
                            set_cell_normal(lc, f"{sug_i + 1}. {s['title']}:\n{ft}")
                        if sug_i < len(qi_sug_upd):
                            u = qi_sug_upd[sug_i]
                            ut = u.get("update_text", INSUFFICIENT)
                            if ut == INSUFFICIENT:
                                set_cell_insufficient(rc)
                            else:
                                set_cell_normal(rc, f"{sug_i + 1}. {u['title']}:\n{ut}")
                        sug_i += 1

    # --- Other domain summary tables ---
    for od in data.get("other_domains", []):
        dkey = get_domain_key(od["domain_name"])
        if not dkey:
            continue
        tidx = DOMAIN_TO_SUMMARY_TABLE.get(dkey)
        if tidx is None or tidx >= len(tables):
            continue

        table = tables[tidx]
        qi_recs = od.get("qi_recommendations", [])
        qi_rec_upd = od.get("qi_recommendation_updates", [])
        qi_sugs = od.get("qi_suggestions", [])
        qi_sug_upd = od.get("qi_suggestion_updates", [])

        in_rec = False
        in_sug = False
        rec_i = 0
        sug_i = 0

        for row in table.rows:
            row_text = " ".join(c.text for c in row.cells)

            if "Evaluation and Comment" in row_text and ("Add comments" in row_text or "no significant changes" in row_text.lower()):
                eval_cell = row.cells[0]
                ev = od.get("evaluation_and_comment", INSUFFICIENT)
                if ev == INSUFFICIENT:
                    set_cell_insufficient(eval_cell)
                else:
                    set_cell_normal(eval_cell, "Evaluation and Comment:\n\n" + ev)
                continue

            if "Quality Improvement Recommendations at" in row_text:
                in_rec = True; in_sug = False; continue
            if "Quality Improvement Suggestions at" in row_text:
                in_sug = True; in_rec = False; continue
            if "Outcome:" in row_text:
                in_rec = False; in_sug = False
                for cell in row.cells:
                    if "Outcome:" in cell.text:
                        set_cell_normal(cell, f"Outcome: {od.get('outcome', 'PASS')}")
                    elif "Rating" in cell.text:
                        set_cell_normal(cell, f"Interim Review Rating: {od.get('rating', 'Good')}")
                continue

            if in_rec and len(row.cells) >= 2:
                lc = row.cells[0]
                rc = row.cells[-1]
                if lc.text.strip() in ["", "1.", "2.", "3."]:
                    if rec_i < len(qi_recs):
                        r = qi_recs[rec_i]
                        ft = r.get("full_text", INSUFFICIENT)
                        if ft == INSUFFICIENT:
                            set_cell_insufficient(lc)
                        else:
                            set_cell_normal(lc, f"{rec_i + 1}. {r['title']}:\n{ft}")
                        if rec_i < len(qi_rec_upd):
                            u = qi_rec_upd[rec_i]
                            ut = u.get("update_text", INSUFFICIENT)
                            if ut == INSUFFICIENT:
                                set_cell_insufficient(rc)
                            else:
                                set_cell_normal(rc, f"{rec_i + 1}. {u['title']}:\n{ut}")
                        rec_i += 1

            if in_sug and len(row.cells) >= 2:
                lc = row.cells[0]
                rc = row.cells[-1]
                if lc.text.strip() in ["", "1.", "2.", "3."]:
                    if sug_i < len(qi_sugs):
                        s = qi_sugs[sug_i]
                        ft = s.get("full_text", INSUFFICIENT)
                        if ft == INSUFFICIENT:
                            set_cell_insufficient(lc)
                        else:
                            set_cell_normal(lc, f"{sug_i + 1}. {s['title']}:\n{ft}")
                        if sug_i < len(qi_sug_upd):
                            u = qi_sug_upd[sug_i]
                            ut = u.get("update_text", INSUFFICIENT)
                            if ut == INSUFFICIENT:
                                set_cell_insufficient(rc)
                            else:
                                set_cell_normal(rc, f"{sug_i + 1}. {u['title']}:\n{ut}")
                        sug_i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/test")
def test_api():
    try:
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=20,
            messages=[{"role": "user", "content": "Say hello in one word"}]
        )
        return jsonify({"status": "ok", "response": msg.content[0].text})
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route("/")
def index():
    return render_template("index.html", domains=DOMAINS)


@app.route("/generate", methods=["POST"])
def generate():
    try:
        pdf_file = request.files.get("reaccreditation_report")
        transcript = request.form.get("transcript", "").strip()
        focus_1 = request.form.get("focus_domain_1")
        focus_2 = request.form.get("focus_domain_2")

        if not pdf_file or not transcript or not focus_1 or not focus_2:
            return jsonify({"error": "All fields are required."}), 400
        if focus_1 == focus_2:
            return jsonify({"error": "Please select two different domains."}), 400

        pdf_text = extract_pdf_text(pdf_file)
        data = get_claude_content(pdf_text, transcript, focus_1, focus_2)

        template_path = os.path.join(os.path.dirname(__file__), "template", "interim_review_template.docx")
        doc_buffer = populate_document(template_path, data)

        unit_name = data.get("unit_name", "Unit").replace(" ", "_").replace(",", "").replace(".", "")
        filename = f"Interim_Review_{unit_name}.docx"

        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except json.JSONDecodeError as e:
        return jsonify({"error": f"Failed to parse AI response: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
